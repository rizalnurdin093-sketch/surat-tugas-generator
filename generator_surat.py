"""
Generator surat tugas .docx dari template (struktur 5 tabel — template baru).

Template (templates_surat/template_surat.docx) punya struktur:
- Paragraph: p4="SURAT TUGAS", p5="Nomor: ...", p33="Daftar Peserta", p34=judul daftar
- Tabel 0 (2x3): Nama/NIK, Jabatan  -> identitas pejabat TTD
- Tabel 1 (3x3): Hari/Tanggal(r0c2), Judul(r1c2), Tempat(r2c2)
- Tabel 2 (4x1): Blok TTD (Ditetapkan di / Pada tanggal)
- Tabel 3 (3x3): Lampiran 1 - Nomor(r1c2), Tanggal(r2c2)
- Tabel 4 (33x4): Daftar peserta - NO, NAMA, NIK, DIVISION

Logika:
- Salin template -> isi field Hari/Tanggal, Judul, Tempat (Tabel 1)
- Isi/update tanggal TTD (Tabel 2) -> bulan sekarang
- Isi nomor surat (paragraf p5 + Tabel 3 r1c2)
- Isi tanggal Lampiran (Tabel 3 r2c2)
- Auto-skalasi + isi tabel peserta (Tabel 4)
- Update judul daftar peserta (p34)
- Set margin sesuai spesifikasi (kiri 3cm, kanan/atas/bawah 2.5cm)
"""

import os
import shutil
import re
import datetime
from copy import deepcopy
from docx import Document
from docx.shared import Cm, Pt

# ============================================================
# Helper: mapping bulan Inggris -> Indonesia
# ============================================================
_BULAN_MAP = {
    "January": "Januari",
    "February": "Februari",
    "March": "Maret",
    "April": "April",
    "May": "Mei",
    "June": "Juni",
    "July": "Juli",
    "August": "Agustus",
    "September": "September",
    "October": "Oktober",
    "November": "November",
    "December": "Desember",
}


def _bulan_indonesia(bulan_eng):
    return _BULAN_MAP.get(bulan_eng, bulan_eng)


TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "templates_surat", "template_surat.docx")

# Indeks struktur template baru (diverifikasi dari template_baru.docx)
P_NOMOR_SURAT = 5          # "Nomor : /Tbk/..."
P_JUDUL_DAFTAR = 34        # judul di bawah "Daftar Peserta"

TABLE_IDENTITAS = 0        # Nama/NIK, Jabatan
TABLE_FORM = 1             # Hari/Tanggal, Judul, Tempat
TABLE_TTD = 2              # blok TTD
TABLE_LAMPIRAN = 3         # Nomor + Tanggal lampiran
TABLE_PESERTA = 4          # daftar peserta

# Margin (cm) sesuai permintaan user: kiri 3, kanan/atas/bawah 2.5
MARGIN_TOP_CM = 2.5
MARGIN_BOTTOM_CM = 2.5
MARGIN_RIGHT_CM = 2.5
MARGIN_LEFT_CM = 3.0

FONT_TABEL = "Arial"
FONT_TABEL_SIZE = 10  # pt


# ============================================================
# Helper fungsi (dipertahankan dari versi lama, s/d adapt struktur)
# ============================================================
def _clean(text):
    if text is None:
        return ""
    return " ".join(str(text).replace("\xa0", " ").split())


def _replace_in_paragraph(paragraph, search, new_value):
    """Ganti 'search' -> 'new_value' di seluruh run paragraph."""
    for run in paragraph.runs:
        if search in run.text:
            run.text = run.text.replace(search, new_value)
            return True
    return False


def _set_cell_first_run(cell, new_value, clear_other_paragraphs=True):
    if clear_other_paragraphs:
        for p in cell.paragraphs[1:]:
            p._element.getparent().remove(p._element)
    first_p = cell.paragraphs[0]
    if first_p.runs:
        first_p.runs[0].text = new_value
        for run in first_p.runs[1:]:
            run.text = ""
    else:
        first_p.add_run(new_value)


def _set_cell_from_query(cell, search, new_value):
    """
    Di dalam satu cell, ganti substring 'search' -> 'new_value'
    pada run yang mengandungnya (untuk cell dengan teks campuran,
    misal 'Batch 11: ...' diganti tanggal).
    """
    for p in cell.paragraphs:
        for run in p.runs:
            if search in run.text:
                run.text = run.text.replace(search, new_value)
    # Fallback: cell tanpa run tersedia -> set cell first run
    return


def _remove_table_rows(table, n_rows_to_remove):
    tbl = table._tbl
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    rows = list(tbl.findall(ns + "tr"))
    for row in rows[-n_rows_to_remove:]:
        tbl.remove(row)


def _add_row_like(table, template_row):
    import copy
    from docx.oxml.ns import qn
    from docx.table import _Row

    new_tr = copy.deepcopy(template_row._tr)
    table._tbl.append(new_tr)
    new_row = _Row(new_tr, table)
    for cell in new_row.cells:
        _clear_cell_text(cell)
    return new_row


def _clear_cell_text(cell):
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    first = cell.paragraphs[0]
    for run in list(first.runs):
        run.text = ""


def _apply_cell_font(cell):
    from docx.shared import Pt
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = FONT_TABEL
            run.font.size = Pt(FONT_TABEL_SIZE)


def _fill_peserta_row(row, no, nama, nik, divisi):
    cells = row.cells
    if len(cells) < 4:
        return
    values = [str(no), _clean(nama), _clean(nik), _clean(divisi)]
    for cell, val in zip(cells, values):
        for p in cell.paragraphs[1:]:
            p._element.getparent().remove(p._element)
        first_p = cell.paragraphs[0]
        if first_p.runs:
            first_p.runs[0].text = val
            for run in first_p.runs[1:]:
                run.text = ""
        else:
            first_p.add_run(val)
        _apply_cell_font(cell)


def _apply_margins(doc):
    """Set margin dokumen sesuai spesifikasi."""
    for section in doc.sections:
        section.top_margin = Cm(MARGIN_TOP_CM)
        section.bottom_margin = Cm(MARGIN_BOTTOM_CM)
        section.left_margin = Cm(MARGIN_LEFT_CM)
        section.right_margin = Cm(MARGIN_RIGHT_CM)


# ============================================================
# Fungsi utama
# ============================================================
def generate_surat(
    judul: str,
    tempat: str,
    hari_tanggal: str,
    peserta: list,
    output_dir: str,
    nomor_angka: str = "",
) -> str:
    """
    Generate 1 file .docx surat tugas (struktur template baru).

    Args:
        judul: judul pelatihan
        tempat: lokasi pelatihan
        hari_tanggal: tanggal pelatihan (sama persis dengan input user)
        peserta: list of dict {'nama', 'nik', 'divisi'}
        nomor_angka: opsional string angka diletakkan di depan '/Tbk/ST-4010/26-S8.7.1'
        output_dir: folder simpan file hasil

    Returns:
        (absolute_path, filename)
    """
    if not peserta:
        raise ValueError("Daftar peserta kosong, tidak bisa generate surat.")

    now = datetime.datetime.now()
    bulan_eng = now.strftime("%B")
    bulan_id = _bulan_indonesia(bulan_eng)

    # Salin template ke output
    safe_judul = re.sub(r"[^\w\s-]", "", judul).strip().replace(" ", "_")[:50]
    output_filename = f"surat_tugas_{safe_judul or 'tanpa_judul'}.docx"
    output_path = os.path.join(output_dir, output_filename)
    shutil.copyfile(TEMPLATE_PATH, output_path)

    doc = Document(output_path)
    _apply_margins(doc)

    # --- 1) Isi formulir: Hari/Tanggal, Judul, Tempat (Tabel 1) ---
    tabel_form = doc.tables[TABLE_FORM]
    # Hari/Tanggal: pakai query replace (ganti semua 'September' lama? tidak —
    #   ganti SELURUH isi cell tanggal dengan hari_tanggal baru)
    sel_hari = tabel_form.rows[0].cells[2]
    _set_cell_first_run(sel_hari, hari_tanggal, clear_other_paragraphs=True)

    sel_judul = tabel_form.rows[1].cells[2]
    _set_cell_first_run(sel_judul, judul, clear_other_paragraphs=True)

    sel_tempat = tabel_form.rows[2].cells[2]
    _set_cell_first_run(sel_tempat, tempat, clear_other_paragraphs=True)

    # --- 2) Nomor surat (paragraf p5 + Tabel 3 lampiran r1c2) ---
    sufix_nomor = "/Tbk/ST-4010/26-S8.7.1"
    if nomor_angka and nomor_angka.strip():
        nomor_lengkap = f"{nomor_angka.strip()} {sufix_nomor}"
        # p5: sufix terpecah di banyak run, jadi rebuild seluruh paragraf
        if P_NOMOR_SURAT < len(doc.paragraphs):
            p_nomor = doc.paragraphs[P_NOMOR_SURAT]
            # Ganti seluruh teks paragraf (pertahankan run pertama utk format)
            if p_nomor.runs:
                # simpan run pertama utk format, isi run[0] = nilai baru, kosongkan sisanya
                p_nomor.runs[0].text = f"Nomor : {nomor_lengkap}"
                for run in p_nomor.runs[1:]:
                    run.text = ""
            else:
                p_nomor.add_run(f"Nomor : {nomor_lengkap}")
        tabel_lampiran = doc.tables[TABLE_LAMPIRAN]
        sel_nomor = tabel_lampiran.rows[1].cells[2]
        _set_cell_from_query(sel_nomor, sufix_nomor, f"{nomor_angka.strip()} {sufix_nomor}")

    # --- 3) Tanggal TTD (Tabel 2) -> bulan sekarang (ID) ---
    tabel_ttd = doc.tables[TABLE_TTD]
    sel_tanggal_ttd = tabel_ttd.rows[0].cells[0]
    for p in sel_tanggal_ttd.paragraphs:
        if "Pada tanggal" in p.text:
            for run in p.runs:
                for b_id_old in _BULAN_MAP.values():
                    if b_id_old in run.text:
                        run.text = run.text.replace(b_id_old, bulan_id)
            break

    # --- 4) Tanggal Lampiran (Tabel 3 r2c2) ---
    tabel_lampiran = doc.tables[TABLE_LAMPIRAN]
    sel_lamp_tanggal = tabel_lampiran.rows[2].cells[2]
    _set_cell_first_run(sel_lamp_tanggal, f"{bulan_id} {now.year}")

    # --- 5) Daftar Peserta (Tabel 4): auto-skalasi + isi ---
    tabel_peserta = doc.tables[TABLE_PESERTA]
    baris_header = 1  # baris header 'NO NAMA NIK DIVISION'
    baris_data_saat_ini = len(tabel_peserta.rows) - baris_header
    baris_data_dibutuhkan = len(peserta)

    if baris_data_dibutuhkan < baris_data_saat_ini:
        _remove_table_rows(tabel_peserta, baris_data_saat_ini - baris_data_dibutuhkan)
    elif baris_data_dibutuhkan > baris_data_saat_ini:
        last_row = tabel_peserta.rows[-1]
        for _ in range(baris_data_dibutuhkan - baris_data_saat_ini):
            _add_row_like(tabel_peserta, last_row)

    for cell in tabel_peserta.rows[0].cells:
        _apply_cell_font(cell)
    for i, p in enumerate(peserta, start=1):
        _fill_peserta_row(tabel_peserta.rows[i], i, p["nama"], p["nik"], p["divisi"])

    # --- 6) Update judul daftar peserta (p34): ganti SELURUH isi paragraf,
    #         karena template berisi judul contoh lama, bukan placeholder 'Judul' ---
    p_judul_lampiran = doc.paragraphs[P_JUDUL_DAFTAR]
    if p_judul_lampiran.runs:
        # Pertahankan run pertama (format teks), isi dengan judul baru
        p_judul_lampiran.runs[0].text = _clean(judul)
        for run in p_judul_lampiran.runs[1:]:
            run.text = ""
    else:
        p_judul_lampiran.add_run(_clean(judul))

    doc.save(output_path)
    return output_path, output_filename
