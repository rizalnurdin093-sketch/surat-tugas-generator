"""Edit statis: ganti penanda tangan di template_surat.docx -> Hikmat Slamet."""
import shutil
from docx import Document

SRC = 'templates_surat/template_surat.docx'
# backup template sebelum edit
shutil.copyfile(SRC, '/tmp/template_surat_backup_pre_hikmat.docx')

doc = Document(SRC)

def set_cell(cell, text):
    """Ganti seluruh isi cell, pertahankan run pertama utk format."""
    # hapus paragraf ekstra, sisakan satu
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    first = cell.paragraphs[0]
    if first.runs:
        first.runs[0].text = text
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(text)

# --- Tabel 0: identitas pejabat ---
t0 = doc.tables[0]
set_cell(t0.rows[0].cells[2], "Hikmat Slamet / 20080624")          # Nama / NIK
set_cell(t0.rows[1].cells[2], "Division Head of Human Capital")      # Jabatan

# --- Tabel 2: blok TTD ---
t2 = doc.tables[2]
set_cell(t2.rows[2].cells[0], "Division Head of Human Capital")      # jabatan TTD
set_cell(t2.rows[3].cells[0], "HIKMAT SLAMET\nNIK. 20080624")        # nama + NIK TTD

doc.save(SRC)
print("Template diperbarui -> Hikmat Slamet")

# verifikasi
d = Document(SRC)
print("Tabel0 r0c2:", repr(d.tables[0].rows[0].cells[2].text))
print("Tabel0 r1c2:", repr(d.tables[0].rows[1].cells[2].text))
print("Tabel2 r2c0:", repr(d.tables[2].rows[2].cells[0].text))
print("Tabel2 r3c0:", repr(d.tables[2].rows[3].cells[0].text))
