"""Self-test generator_surat dengan template BARU (tanpa mail merge)."""
import os
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
sys.path.insert(0, "/home/ubuntu/.local/lib/python3.12/site-packages")

from parser_excel import parse_excel, sort_az
from generator_surat import generate_surat
from docx import Document

peserta = sort_az(parse_excel("uploads/contoh_peserta.xlsx"))[:4]
print(f"Peserta test: {len(peserta)}")

out = "/tmp/surat-test2"
os.makedirs(out, exist_ok=True)
path, fname = generate_surat(
    judul="Pengelolaan SDM IHT",
    tempat="Ruang Meeting HC",
    hari_tanggal="Kamis - Jumat / 20 Agustus 2026 - 21 Agustus 2026",
    tanggal_surat="Agustus 2026",
    peserta=peserta,
    output_dir=out,
)
print(f"Generated: {path}\n")

doc = Document(path)

# ── FIELD (P10/P11/P12) ──
print("=== BODY FIELD ===")
for i,label in [(10,"Hari/Tanggal"),(11,"Judul"),(12,"Tempat")]:
    p = doc.paragraphs[i]
    runs = [(r.text, bool(r.bold)) for r in p.runs]
    print(f"P{i} {label}: {repr(p.text)} runs={runs}")

print("\n=== NOMOR SURAT (P1) — harus tetap /Tbk/... ===")
print("  P1:", repr(doc.paragraphs[1].text))

print("\n=== TABEL 1 (Lampiran) — nomor tetap, tanggal bold+spasi ===")
t1 = doc.tables[1]
print("  row1 (Nomor):", [repr(c.text) for c in t1.rows[1].cells])
r2 = t1.rows[2].cells[2]
print("  row2 cell2 (Tanggal):", repr(r2.text), "| bold:", r2.paragraphs[0].runs[0].bold)

print("\n=== TABEL 0 (TTD) row0 ===")
print("  row0:", [repr(c.text) for c in doc.tables[0].rows[0].cells])

print("\n=== P35 (Judul lampiran) ===")
p35 = doc.paragraphs[35]
print("  P35:", repr(p35.text), "| bold:", [bool(r.bold) for r in p35.runs])

print("\n=== TABEL PESERTA ===")
t2 = doc.tables[2]
print("  rows:", len(t2.rows), "(4 peserta + header = 5)")

# ── Asserts ──
def p_contains(i, txt): return txt in doc.paragraphs[i].text
assert p_contains(10, "Kamis - Jumat"), "Hari/Tanggal terisi"
assert p_contains(11, "Pengelolaan SDM IHT"), "Judul terisi"
assert p_contains(12, "Ruang Meeting HC"), "Tempat terisi"
assert p_contains(35, "Pengelolaan SDM IHT"), "Judul lampiran terisi"

# Nomor surat & nomor lampiran dipertahankan
assert "/Tbk/ST-4010/26-S8.7.1" in doc.paragraphs[1].text, "P1 tetap berisi /Tbk/..."
assert "/Tbk/ST-4010/26-S8.7.1" in t1.rows[1].cells[2].text, "T1 nomor tetap"

# Tanggal lampiran bold + spasi depan
assert t1.rows[2].cells[2].paragraphs[0].runs[0].bold, "Tanggal lampiran harus bold"
assert t1.rows[2].cells[2].text.startswith(" "), "Tanggal lampiran ada spasi depan"

# P10 nilai harus bold (run placeholder bold)
p10_bold = [b for _,b in [(r.text,bool(r.bold)) for r in doc.paragraphs[10].runs]]
assert any(b and "2026" in r.text for r in doc.paragraphs[10].runs for b in [bool(r.bold)]), "Hari/Tanggal nilai bold"

# Jumlah tabel peserta
assert len(t2.rows) == 5, f"rows={len(t2.rows)}"

print("\n✅ SELF-TEST PASSED — semua koreksi diterapkan & format terjaga")
